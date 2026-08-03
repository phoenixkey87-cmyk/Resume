"""
Resume Builder for LaKeysha Strickland
Generates a professionally designed Word (.docx) and PDF version of the resume.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from fpdf import FPDF
import os

# --- Hyperlinks ---
LINKEDIN_URL = "https://www.linkedin.com/in/lakeysha-strickland/"
GITHUB_PROJECT_URL = "https://phoenixkey87-cmyk.github.io/Keysha_40th_Vietnam/"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# --- Color Palette ---
NAVY = RGBColor(0x1B, 0x3A, 0x5C)  # Deep navy for headings
ACCENT = RGBColor(0x2E, 0x86, 0xAB)  # Teal accent
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)  # Body text
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)  # Background accents

# --- Resume Data ---
NAME = "LaKeysha Strickland, MBA, CSM, CSPO"
TITLE = "SYSTEMS ENGINEER II | AI-ASSISTED SOLUTIONS | ENTERPRISE AUTOMATION | TECHNICAL PRODUCT DELIVERY"
CONTACT = "St. Louis, MO  |  314-307-4921  |  keystrickland@charter.net  |  LinkedIn  |  Live Kiro/GitHub Project"

SUMMARY = (
    "Systems Engineer II with 12+ years of progressive experience spanning enterprise network operations, "
    "change governance, AI-assisted solution development, process automation, and technical training. "
    "Designed and launched an enterprise Deconfliction Dashboard and Portal using Amazon Kiro and Microsoft "
    "Power Apps that proactively identifies maintenance conflicts and helps avoid at least 10 potential outages "
    "on a typical day, including 19 in a single day. Coordinates high-volume change activity across 20+ "
    "engineering and operations teams while translating complex technical needs into practical tools, "
    "documented processes, training, and measurable operational improvements."
)

IMPACT_METRICS = [
    ("10\u201319", "potential outages avoided daily"),
    ("20+", "carrier tickets coordinated daily"),
    ("50+", "enterprise changes handled weekly"),
    ("20+", "engineering teams supported"),
    ("6", "SOPs updated and standardized"),
    ("Multiple", "employees trained across processes"),
    ("Enterprise", "dashboard built from concept to launch"),
    ("Live", "customer-facing sites deployed via GitHub"),
]

CORE_EXPERTISE = [
    "Amazon Kiro", "Amazon Q", "Microsoft Power Apps", "Power Automate", "AWS", "Azure",
    "GitHub/GitHub Pages", "AI-Assisted Development", "Systems Engineering",
    "Enterprise Network Operations", "Change Management", "Risk & Impact Analysis",
    "Technical Product Delivery", "Agile/Scrum", "ITIL", "Jira", "BMC Remedy",
    "SharePoint", "UAT", "SOP Development", "Technical Training", "Stakeholder Management"
]

EXPERIENCE = [
    {
        "title": "Systems Engineer II",
        "company": "Spectrum / Charter Communications",
        "dates": "Mar 2026 \u2013 Present",
        "bullets": [
            "Designed and deployed an enterprise-wide Deconfliction Dashboard and Portal using Amazon Kiro and Microsoft Power Apps to identify overlapping maintenance activity before implementation.",
            "Help avoid at least 10 potential customer-impacting outages on a typical day through proactive conflict detection, including 19 potential outages identified and prevented in a single day.",
            "Coordinate approximately 50 enterprise change tickets per week and roughly 20 third-party carrier tickets per day for providers including Comcast, Lumen, Netflix, and other telecommunications and content partners.",
            "Collaborate with 20+ engineering and operations teams nationwide to validate impacts, correct ticket data, enforce approved Methods of Procedure (MOPs), and reduce operational risk.",
            "Updated and standardized six SOPs to strengthen governance, improve consistency, and clarify execution requirements across engineering workflows.",
            "Train employees on impact corrections, carrier ticket creation, correct MOP selection, change processes, and internal tools; also create video guides and technical documentation to support adoption.",
            "Use Amazon Q for code support, documentation, testing, workflow automation, and process improvement; represent the team in cloud-related meetings and communicate updates to stakeholders.",
        ],
    },
    {
        "title": "Change Control Analyst II",
        "company": "Spectrum / Charter Communications",
        "dates": "Mar 2019 \u2013 Mar 2026",
        "bullets": [
            "Led analytic oversight of scheduled network maintenance, identifying resource and cross-service collisions across a converged enterprise network.",
            "Served as a primary point of contact for third-party maintenance activity and coordinated complex changes among engineering, operations, project teams, and external providers.",
            "Prepared operational reporting, analyzed maintenance trends, and recommended standards and process improvements that strengthened compliance and network reliability.",
            "Developed change support documentation, manuals, procedures, and training while performing impact assessments and risk reviews in alignment with ITIL practices.",
        ],
    },
    {
        "title": "Change Control Analyst",
        "company": "Spectrum / Charter Communications",
        "dates": "Jan 2017 \u2013 Mar 2019",
        "bullets": [
            "Coordinated and facilitated scheduled network maintenance while enforcing change policies and procedures designed to protect the Spectrum network and revenue-generating customers.",
            "Analyzed customer and service impacts, escalated conflicts, and supported carrier-class operating practices across enterprise maintenance activity.",
        ],
    },
    {
        "title": "Internet & Voice Repair Technician",
        "company": "Spectrum / Charter Communications",
        "dates": "Jan 2014 \u2013 Jan 2017",
        "bullets": [
            "Resolved complex internet and voice service issues, handled customer escalations, and consistently supported service-quality and customer-experience goals.",
            "Earned the Spectrum BEST Award for outstanding customer service performance.",
        ],
    },
]

PROJECTS = [
    {
        "title": "Enterprise Deconfliction Dashboard & Portal",
        "tech": "Amazon Kiro, Microsoft Power Apps, AI-Assisted Development",
        "bullets": [
            "Took the solution from business problem and concept through design, development, testing, user adoption, and ongoing enhancement; partnered with a colleague for SQL integration support.",
            "Centralized maintenance visibility so specialists can identify ticket collisions before execution and make informed decisions that reduce outage risk and manual research.",
        ],
    },
    {
        "title": "Phoenix Travel Digital Experiences",
        "tech": "Amazon Kiro, GitHub, GitHub Pages, HTML/CSS",
        "bullets": [
            "Built and deployed live customer-facing travel sites featuring interactive itineraries, RSVP pages, and payment summaries for group travel clients.",
            "Used AI-assisted development to turn business requirements into polished web experiences, then managed source control and production publishing through GitHub Pages.",
            "View live project: https://phoenixkey87-cmyk.github.io/Keysha_40th_Vietnam/",
        ],
    },
]

EDUCATION = [
    ("Master of Business Administration (MBA)", "Louisiana State University", "2026"),
    ("Bachelor of Science, Business Administration", "Lindenwood University", "2021"),
]

CERTIFICATIONS = [
    "Certified ScrumMaster (CSM), Scrum Alliance, 2022",
    "Certified Scrum Product Owner (CSPO), Scrum Alliance, 2023",
    "NYU Project Management Certificate",
]

AWARDS = [
    "Spectrum BEST Award \u2013 recognized for outstanding customer service performance.",
    "Active in NAMIC and WICT professional-development initiatives; Make-A-Wish volunteer.",
    "Selected by a director for consideration for the WICT Women to Watch recognition.",
]


# =============================================================================
# WORD DOCUMENT GENERATION
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, url, text, font_size=Pt(9), color=RGBColor(0x2E, 0x86, 0xAB)):
    """Add a clickable hyperlink to a paragraph in a Word document."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'</w:hyperlink>'
    )

    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr>'
        f'<w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="{color}"/>'
        f'<w:u w:val="single"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'</w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, text):
    """Add a styled section heading with a bottom border."""
    para = doc.add_paragraph()
    para.space_before = Pt(14)
    para.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    # Add bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="2E86AB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    para.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
    else:
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT


def build_word_resume():
    """Generate the Word document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- Name ---
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(0)
    run = para.add_run("LaKeysha Strickland")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    # --- Credentials ---
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_before = Pt(0)
    para.space_after = Pt(2)
    run = para.add_run("MBA, CSM, CSPO")
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT

    # --- Title ---
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(4)
    run = para.add_run(TITLE)
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_TEXT
    run.bold = True

    # --- Contact ---
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(8)
    # Build contact line with hyperlinks
    run = para.add_run("St. Louis, MO  |  314-307-4921  |  keystrickland@charter.net  |  ")
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_TEXT
    add_hyperlink(para, LINKEDIN_URL, "LinkedIn", font_size=Pt(9))
    run = para.add_run("  |  ")
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_TEXT
    add_hyperlink(para, GITHUB_PROJECT_URL, "Live Kiro/GitHub Project", font_size=Pt(9))

    # --- Professional Summary ---
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    para = doc.add_paragraph()
    para.space_after = Pt(6)
    run = para.add_run(SUMMARY)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT

    # --- Selected Impact (table) ---
    add_section_heading(doc, "SELECTED IMPACT")
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    metrics_row1 = IMPACT_METRICS[:4]
    metrics_row2 = IMPACT_METRICS[4:]

    for i, (value, label) in enumerate(metrics_row1):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = ACCENT
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = DARK_TEXT
        set_cell_shading(cell, "F0F7FA")

    for i, (value, label) in enumerate(metrics_row2):
        cell = table.rows[1].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = ACCENT
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = DARK_TEXT
        set_cell_shading(cell, "F0F7FA")

    # --- Core Expertise ---
    add_section_heading(doc, "CORE EXPERTISE")
    para = doc.add_paragraph()
    para.space_after = Pt(6)
    run = para.add_run("  \u2022  ".join(CORE_EXPERTISE))
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_TEXT

    # --- Professional Experience ---
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    for job in EXPERIENCE:
        # Job title + dates
        para = doc.add_paragraph()
        para.space_before = Pt(8)
        para.space_after = Pt(0)
        run = para.add_run(job["title"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run = para.add_run(f"  |  {job['company']}")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

        para2 = doc.add_paragraph()
        para2.space_before = Pt(0)
        para2.space_after = Pt(4)
        run = para2.add_run(job["dates"])
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = ACCENT

        for bullet in job["bullets"]:
            add_bullet(doc, bullet)

    # --- Selected Technical Projects ---
    add_section_heading(doc, "SELECTED TECHNICAL PROJECTS")
    for proj in PROJECTS:
        para = doc.add_paragraph()
        para.space_before = Pt(8)
        para.space_after = Pt(2)
        run = para.add_run(proj["title"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run = para.add_run(f"  |  {proj['tech']}")
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_TEXT

        for bullet in proj["bullets"]:
            add_bullet(doc, bullet)

    # --- Education ---
    add_section_heading(doc, "EDUCATION")
    for degree, school, year in EDUCATION:
        para = doc.add_paragraph()
        para.space_after = Pt(2)
        run = para.add_run(degree)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
        run = para.add_run(f"  |  {school}  |  {year}")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

    # Dean's List note
    para = doc.add_paragraph()
    para.space_after = Pt(2)
    run = para.add_run("    Dean\u2019s List")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_TEXT

    # --- Certifications ---
    add_section_heading(doc, "CERTIFICATIONS & PROFESSIONAL DEVELOPMENT")
    para = doc.add_paragraph()
    para.space_after = Pt(4)
    run = para.add_run("  \u2022  ".join(CERTIFICATIONS))
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT

    # --- Awards ---
    add_section_heading(doc, "AWARDS & COMMUNITY LEADERSHIP")
    for award in AWARDS:
        add_bullet(doc, award)

    # Save
    docx_path = os.path.join(OUTPUT_DIR, "LaKeysha_Strickland_Resume.docx")
    doc.save(docx_path)
    print(f"[OK] Word document saved: {docx_path}")
    return docx_path


# =============================================================================
# PDF GENERATION
# =============================================================================

def sanitize_for_pdf(text):
    """Replace Unicode characters that core fonts can't render."""
    replacements = {
        "\u2013": "-",   # en-dash
        "\u2014": "-",   # em-dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2022": "-",   # bullet (we'll handle bullets separately)
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text


class ResumePDF(FPDF):
    """Custom PDF class for resume formatting."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def safe_cell(self, w, h, text, **kwargs):
        """Cell with sanitized text."""
        self.cell(w, h, sanitize_for_pdf(text), **kwargs)

    def safe_multi_cell(self, w, h, text, **kwargs):
        """Multi_cell with sanitized text."""
        self.multi_cell(w, h, sanitize_for_pdf(text), **kwargs)

    def header_section(self):
        """Render the name and contact header."""
        # Name
        self.set_font("Helvetica", "B", 22)
        self.set_text_color(27, 58, 92)  # Navy
        self.safe_cell(0, 10, "LaKeysha Strickland", align="C", new_x="LMARGIN", new_y="NEXT")

        # Credentials
        self.set_font("Helvetica", "", 11)
        self.set_text_color(46, 134, 171)  # Accent
        self.safe_cell(0, 6, "MBA, CSM, CSPO", align="C", new_x="LMARGIN", new_y="NEXT")

        # Title
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(45, 45, 45)
        self.safe_cell(0, 6, TITLE, align="C", new_x="LMARGIN", new_y="NEXT")

        # Contact line with hyperlinks
        self.set_font("Helvetica", "", 8)
        self.set_text_color(45, 45, 45)
        contact_prefix = "St. Louis, MO  |  314-307-4921  |  keystrickland@charter.net  |  "
        contact_sep = "  |  "
        linkedin_text = "LinkedIn"
        github_text = "Live Kiro/GitHub Project"

        # Calculate total width to center it
        prefix_w = self.get_string_width(contact_prefix)
        linkedin_w = self.get_string_width(linkedin_text)
        sep_w = self.get_string_width(contact_sep)
        github_w = self.get_string_width(github_text)
        total_w = prefix_w + linkedin_w + sep_w + github_w
        start_x = (self.w - total_w) / 2

        y = self.get_y()
        self.set_xy(start_x, y)
        self.set_text_color(45, 45, 45)
        self.cell(prefix_w, 5, sanitize_for_pdf(contact_prefix))

        # LinkedIn link
        self.set_text_color(46, 134, 171)
        self.cell(linkedin_w, 5, linkedin_text, link=LINKEDIN_URL)

        self.set_text_color(45, 45, 45)
        self.cell(sep_w, 5, contact_sep)

        # GitHub Project link
        self.set_text_color(46, 134, 171)
        self.cell(github_w, 5, github_text, link=GITHUB_PROJECT_URL)

        self.ln(8)

    def section_heading(self, text):
        """Render a section heading with bottom line."""
        self.ln(4)
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(27, 58, 92)
        self.safe_cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        # Draw accent line
        self.set_draw_color(46, 134, 171)
        self.set_line_width(0.5)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body_text(self, text):
        """Render body paragraph text."""
        self.set_font("Helvetica", "", 9)
        self.set_text_color(45, 45, 45)
        self.safe_multi_cell(0, 4.5, text)
        self.ln(2)

    def bullet_point(self, text):
        """Render a bullet point."""
        self.set_font("Helvetica", "", 9)
        self.set_text_color(45, 45, 45)
        x = self.get_x()
        self.set_x(x + 4)
        self.safe_multi_cell(0, 4.5, f"-  {text}")
        self.ln(1)

    def job_header(self, title, company, dates):
        """Render a job title line."""
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(27, 58, 92)
        self.safe_cell(0, 6, f"{title}  |  {company}", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(46, 134, 171)
        self.safe_cell(0, 5, dates, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)


def build_pdf_resume():
    """Generate the PDF resume."""
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 15, 18)

    pdf.header_section()

    # Professional Summary
    pdf.section_heading("PROFESSIONAL SUMMARY")
    pdf.body_text(SUMMARY)

    # Selected Impact
    pdf.section_heading("SELECTED IMPACT")
    # Render as a compact grid
    col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / 4
    start_x = pdf.get_x()
    start_y = pdf.get_y()

    for row_idx in range(2):
        row_metrics = IMPACT_METRICS[row_idx * 4:(row_idx + 1) * 4]
        for col_idx, (value, label) in enumerate(row_metrics):
            x = start_x + col_idx * col_w
            y = start_y + row_idx * 14
            # Value
            pdf.set_xy(x, y)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(46, 134, 171)
            pdf.safe_cell(col_w, 5, value, align="C")
            # Label
            pdf.set_xy(x, y + 5)
            pdf.set_font("Helvetica", "", 7)
            pdf.set_text_color(80, 80, 80)
            pdf.safe_cell(col_w, 4, label, align="C")

    pdf.set_y(start_y + 32)

    # Core Expertise
    pdf.section_heading("CORE EXPERTISE")
    expertise_text = "  |  ".join(CORE_EXPERTISE)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(45, 45, 45)
    pdf.safe_multi_cell(0, 4, expertise_text)
    pdf.ln(2)

    # Professional Experience
    pdf.section_heading("PROFESSIONAL EXPERIENCE")
    for job in EXPERIENCE:
        pdf.job_header(job["title"], job["company"], job["dates"])
        for bullet in job["bullets"]:
            pdf.bullet_point(bullet)
        pdf.ln(2)

    # Selected Technical Projects
    pdf.section_heading("SELECTED TECHNICAL PROJECTS")
    for proj in PROJECTS:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(27, 58, 92)
        pdf.safe_cell(0, 6, f"{proj['title']}  |  {proj['tech']}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
        for bullet in proj["bullets"]:
            pdf.bullet_point(bullet)
        pdf.ln(2)

    # Education
    pdf.section_heading("EDUCATION")
    for degree, school, year in EDUCATION:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(45, 45, 45)
        pdf.safe_cell(0, 5, f"{degree}  |  {school}  |  {year}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(80, 80, 80)
    pdf.safe_cell(0, 5, "    Dean's List", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Certifications
    pdf.section_heading("CERTIFICATIONS & PROFESSIONAL DEVELOPMENT")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(45, 45, 45)
    pdf.safe_multi_cell(0, 4.5, "  |  ".join(CERTIFICATIONS))
    pdf.ln(2)

    # Awards
    pdf.section_heading("AWARDS & COMMUNITY LEADERSHIP")
    for award in AWARDS:
        pdf.bullet_point(award)

    # Save
    pdf_path = os.path.join(OUTPUT_DIR, "LaKeysha_Strickland_Resume.pdf")
    pdf.output(pdf_path)
    print(f"[OK] PDF saved: {pdf_path}")
    return pdf_path


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    print("Building resume documents...")
    print("-" * 50)
    build_word_resume()
    build_pdf_resume()
    print("-" * 50)
    print("Done! Files are in the output/ folder.")
