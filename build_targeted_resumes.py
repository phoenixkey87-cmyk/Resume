"""
Targeted Resume Builder for LaKeysha Strickland
Generates 3 tailored versions: AI Automation Engineer, Solutions Engineer, TPM
Each version gets HTML, PDF, and Word (.docx) output.
"""

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# --- Shared Data ---
NAME = "LaKeysha Strickland"
CREDENTIALS = "MBA, CSM, CSPO, AWS AI Practitioner"
PHONE = "314-307-4921"
EMAIL = "keystrickland@charter.net"
LOCATION = "St. Louis, MO"
LINKEDIN = "https://www.linkedin.com/in/lakeysha-strickland/"
GITHUB = "https://github.com/phoenixkey87-cmyk"

EDUCATION = [
    ("Master of Business Administration (MBA)", "Louisiana State University", "2026"),
    ("Bachelor of Science, Business Administration", "Lindenwood University - Dean's List", "2021"),
]

CERTIFICATIONS = [
    "AWS Certified AI Practitioner",
    "Certified Scrum Product Owner (CSPO), Scrum Alliance, 2023",
    "Certified ScrumMaster (CSM), Scrum Alliance, 2022",
    "Project Management Certificate, NYU (Pathstream), 2024",
    "Six Sigma: Green Belt",
    "Six Sigma: Yellow Belt",
    "Scrum Fundamentals Certified (SFC)",
]

AWARDS = [
    "Spectrum BEST Award - recognized for outstanding customer service performance.",
    "Deconfliction Dashboard adopted enterprise-wide within weeks of launch.",
    "Selected by a director for consideration for the WICT Women to Watch recognition.",
    "Active in NAMIC and WICT professional-development initiatives; Make-A-Wish volunteer.",
]

# =============================================================================
# VERSION 1: AI AUTOMATION ENGINEER
# =============================================================================
AI_AUTO = {
    "title_line": "AI Automation Engineer | Python | AWS | Enterprise Workflow Automation",
    "summary": (
        "AI-focused automation engineer with 12+ years in enterprise operations, building "
        "production-ready tools with Python, Amazon Kiro, Bedrock, and Microsoft Power Platform. "
        "Architected an enterprise Deconfliction Dashboard preventing 10-19 outages daily and "
        "automated carrier maintenance workflows across 13+ providers, cutting manual processing "
        "from 15 minutes to under 2 minutes per ticket. AWS Certified AI Practitioner combining "
        "hands-on AI-assisted development with deep operational domain expertise."
    ),
    "skills": {
        "AI & Automation": [
            "Python", "Amazon Kiro", "Amazon Q", "Amazon Bedrock", "Claude Code",
            "AI-Assisted Development", "Outlook COM Automation", "Regex", "REST APIs",
            "Workflow Automation", "openpyxl"
        ],
        "Cloud & Platforms": [
            "AWS", "Azure", "Power Apps", "Power Automate", "SharePoint",
            "Microsoft Teams Webhooks", "GitHub / GitHub Pages"
        ],
        "Operations & Data": [
            "Change Management", "ITIL", "Risk & Impact Analysis", "SOP Development",
            "Tableau", "Excel", "BMC Helix", "Jira", "Asana"
        ],
    },
    "experience": [
        {
            "title": "Systems Engineer II - AI-Assisted Automation & Change Management",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2026 - Present",
            "bullets": [
                "Architected an enterprise Deconfliction Dashboard using Amazon Kiro and Power Apps that identifies maintenance conflicts before implementation, preventing 10-19 potential outages daily and saving 2-3 hours of manual research per shift.",
                "Engineered a Python-based carrier maintenance automation system using Outlook COM, regex parsing, and Teams webhooks that detects notifications from 13+ carriers and auto-generates CRQ records -- reducing processing time from ~15 minutes to under 2 minutes per ticket.",
                "Designed impact auto-determination logic that classifies maintenance severity by emergency status, subscriber count, and circuit count, then maps to carrier-specific assignee groups automatically.",
                "Built persistent tracking with duplicate detection and full ticket lifecycle management from email detection through CRQ creation and resolution.",
                "Leverage Amazon Q and Claude Code for rapid prototyping, code generation, testing, and documentation across all automation projects.",
                "Coordinate ~50 enterprise change tickets weekly and ~20 carrier tickets daily with zero administrative errors.",
                "Produced 10+ video guides and technical documentation resources adopted team-wide to accelerate tool adoption.",
            ],
        },
        {
            "title": "Change Control Analyst II",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2019 - Mar 2026",
            "bullets": [
                "Led analytic oversight of 1,000+ scheduled maintenance changes annually across a converged enterprise network serving millions of subscribers.",
                "Prepared weekly and monthly operational reporting and recommended process improvements that reduced repeat escalations.",
                "Authored training materials adopted across 7 engineering verticals, improving CRQ accuracy and reducing revision cycles.",
            ],
        },
        {
            "title": "Change Control Analyst",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Jan 2017 - Mar 2019",
            "bullets": [
                "Coordinated scheduled network maintenance while enforcing change policies protecting revenue-generating customers.",
                "Analyzed customer and service impacts, escalated conflicts, and supported carrier-class operating practices.",
            ],
        },
    ],
    "projects": [
        {
            "title": "Enterprise Deconfliction Dashboard & Portal",
            "tech": "Amazon Kiro, Power Apps, SharePoint, AI-Assisted Development",
            "bullets": [
                "Took the solution from concept through design, development, testing, and enterprise adoption; partnered with a colleague for SQL integration.",
                "Centralized maintenance visibility, eliminating hours of manual research daily and enabling data-driven go/no-go decisions on scheduled maintenance.",
            ],
        },
        {
            "title": "Third-Party Carrier Maintenance Automation",
            "tech": "Python, Outlook COM (pywin32), Teams Webhooks, openpyxl, Amazon Kiro",
            "bullets": [
                "End-to-end workflow: mailbox monitoring, email parsing, Teams alerts, CRQ auto-generation, and persistent tracking with duplicate detection.",
                "Configurable regex patterns for ticket extraction across 13+ carrier formats; impact auto-determination logic for severity classification.",
            ],
        },
        {
            "title": "Phoenix Travel Digital Experiences",
            "tech": "Amazon Kiro, GitHub Pages, HTML/CSS",
            "bullets": [
                "Built and deployed live customer-facing travel sites using AI-assisted development, managing source control and production publishing through GitHub Pages.",
            ],
        },
    ],
}

# =============================================================================
# VERSION 2: SOLUTIONS ENGINEER
# =============================================================================
SOLUTIONS_ENG = {
    "title_line": "Solutions Engineer | Technical Product Delivery | Customer-Facing Systems",
    "summary": (
        "Solutions-oriented engineer with 12+ years translating complex technical requirements "
        "into production-ready tools and customer-facing digital experiences. Combines hands-on "
        "development (Python, Power Platform, AWS) with cross-functional stakeholder management "
        "across 20+ engineering teams. Proven ability to demo, design, and deliver solutions that "
        "prevent outages, automate workflows, and drive measurable adoption -- including an "
        "enterprise dashboard adopted company-wide within weeks of launch."
    ),
    "skills": {
        "Solution Design & Delivery": [
            "Technical Product Delivery", "Requirements Gathering", "Solution Architecture",
            "Stakeholder Management", "UAT", "AI-Assisted Development", "Technical Training",
            "Demo & Presentation"
        ],
        "Development & Platforms": [
            "Python", "Amazon Kiro", "Amazon Q", "Power Apps", "Power Automate",
            "SharePoint", "GitHub / GitHub Pages", "REST APIs", "AWS", "Amazon Bedrock", "Azure"
        ],
        "Operations & Process": [
            "Change Management", "ITIL", "Agile / Scrum", "Risk & Impact Analysis",
            "SOP Development", "BMC Helix", "Jira", "Asana", "Tableau"
        ],
    },
    "experience": [
        {
            "title": "Systems Engineer II - Solution Design & Automation",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2026 - Present",
            "bullets": [
                "Identified a critical operational gap, designed a solution, and delivered an enterprise Deconfliction Dashboard using Amazon Kiro and Power Apps -- adopted company-wide within weeks, preventing 10-19 potential outages daily.",
                "Gathered requirements from 20+ engineering teams, translated them into technical specifications, and built automation tools that eliminated manual workflows and reduced ticket processing from 15 minutes to under 2 minutes.",
                "Serve as technical liaison between Change Management and engineering teams nationwide, validating impacts, demonstrating tools, and driving adoption of new processes.",
                "Produced 10+ training videos and documentation resources, onboarding new users and ensuring consistent tool adoption across the organization.",
                "Coordinate ~50 enterprise change tickets weekly and ~20 carrier tickets daily with zero administrative errors, acting as the team's subject matter expert on carrier processes.",
                "Represent the team in cloud-related meetings, communicate technical updates to non-technical stakeholders, and translate business needs into actionable automation.",
            ],
        },
        {
            "title": "Change Control Analyst II",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2019 - Mar 2026",
            "bullets": [
                "Served as primary point of contact for third-party carrier maintenance, coordinating complex changes among engineering, operations, and external providers with zero missed SLA windows.",
                "Led analytic oversight of 1,000+ maintenance changes annually, presenting findings to leadership and recommending process improvements that reduced escalations.",
                "Authored documentation and training materials adopted across 7 verticals, becoming the go-to resource for change process questions.",
            ],
        },
        {
            "title": "Change Control Analyst / Internet & Voice Repair Technician",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Jan 2014 - Mar 2019",
            "bullets": [
                "Progressed from frontline customer support to change coordination, consistently delivering solutions that improved customer experience and earned the Spectrum BEST Award.",
                "Analyzed customer impacts, escalated conflicts, and supported carrier-class operating practices across enterprise maintenance activity.",
            ],
        },
    ],
    "projects": [
        {
            "title": "Enterprise Deconfliction Dashboard & Portal",
            "tech": "Amazon Kiro, Power Apps, SharePoint, AI-Assisted Development",
            "bullets": [
                "Owned the full solution lifecycle: problem identification, stakeholder interviews, design, build, UAT, launch, and ongoing enhancement.",
                "Delivered a self-service tool that gives 20+ teams instant visibility into maintenance conflicts, replacing a manual multi-step research process.",
            ],
        },
        {
            "title": "Third-Party Carrier Maintenance Automation",
            "tech": "Python, Outlook COM, Teams Webhooks, openpyxl, Amazon Kiro",
            "bullets": [
                "Designed and delivered an automated notification system that gives the team real-time visibility into carrier work, eliminating manual inbox monitoring.",
                "Pre-fills CRQ records with impact classification, reducing handoff time and enabling faster response to provider maintenance.",
            ],
        },
        {
            "title": "Phoenix Travel Digital Experiences",
            "tech": "Amazon Kiro, GitHub Pages, HTML/CSS",
            "bullets": [
                "Built and deployed customer-facing travel sites with interactive itineraries and payment summaries, demonstrating ability to deliver polished end-user experiences.",
            ],
        },
    ],
}

# =============================================================================
# VERSION 3: TECHNICAL PROGRAM MANAGER
# =============================================================================
TPM = {
    "title_line": "Technical Program Manager | Cross-Functional Delivery | Process & Governance",
    "summary": (
        "Technical Program Manager with 12+ years driving cross-functional delivery across "
        "enterprise network operations, change governance, and automation initiatives. Led "
        "programs spanning 20+ engineering teams, coordinating 50+ changes weekly while "
        "standardizing processes, reducing risk, and delivering tools adopted enterprise-wide. "
        "MBA-educated with Scrum (CSM/CSPO), Six Sigma, and NYU PM credentials. Combines "
        "technical fluency (Python, AWS, Power Platform) with proven stakeholder management "
        "and governance leadership."
    ),
    "skills": {
        "Program & Project Management": [
            "Cross-Functional Leadership", "Stakeholder Management", "Agile / Scrum",
            "ITIL Framework", "Six Sigma", "Risk & Impact Analysis", "Governance",
            "SOP Development", "Technical Training", "Change Management"
        ],
        "Technical Delivery": [
            "Python", "Amazon Kiro", "Amazon Q", "Power Apps", "Power Automate",
            "AWS", "Amazon Bedrock", "Azure", "SharePoint", "GitHub"
        ],
        "Tools & Reporting": [
            "Jira", "Asana", "BMC Helix", "Tableau", "Excel",
            "Technical Documentation", "Executive Reporting", "UAT"
        ],
    },
    "experience": [
        {
            "title": "Systems Engineer II - Program Delivery & Automation",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2026 - Present",
            "bullets": [
                "Led the end-to-end delivery of an enterprise Deconfliction Dashboard -- from problem identification through stakeholder alignment, development, UAT, launch, and adoption -- now used daily by the Change Management team to prevent 10-19 potential outages.",
                "Drove the carrier maintenance automation initiative from concept to production, coordinating requirements, development, testing, and rollout across the team.",
                "Manage the coordination of ~50 enterprise change tickets weekly and ~20 carrier tickets daily across 20+ engineering and operations teams nationwide.",
                "Standardized 6 SOPs, reducing onboarding time for new analysts and improving first-pass accuracy across the department.",
                "Produced 10+ training resources and led knowledge-transfer sessions that accelerated adoption of new tools and processes team-wide.",
                "Represent the team in cloud-related program meetings, communicate status and roadmap to stakeholders, and align priorities with organizational objectives.",
            ],
        },
        {
            "title": "Change Control Analyst II - Process Governance & Coordination",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Mar 2019 - Mar 2026",
            "bullets": [
                "Oversaw governance for 1,000+ scheduled maintenance changes annually, ensuring compliance with ITIL practices and protecting a network serving millions of subscribers.",
                "Served as primary coordinator for third-party carrier maintenance, managing relationships with external providers and maintaining zero missed SLA windows.",
                "Prepared weekly/monthly executive reporting, analyzed trends, and recommended process improvements that reduced repeat escalations and strengthened reliability.",
                "Authored process documentation and training adopted across 7 engineering verticals, establishing consistent standards organization-wide.",
            ],
        },
        {
            "title": "Change Control Analyst / Internet & Voice Repair Technician",
            "company": "Spectrum (Charter Communications)  |  St. Louis, MO",
            "dates": "Jan 2014 - Mar 2019",
            "bullets": [
                "Progressed through 3 roles in 5 years based on consistent delivery, process improvement contributions, and leadership recognition (Spectrum BEST Award).",
                "Coordinated network maintenance, enforced change policies, and analyzed impacts across enterprise services.",
            ],
        },
    ],
    "projects": [
        {
            "title": "Enterprise Deconfliction Dashboard & Portal",
            "tech": "Amazon Kiro, Power Apps, SharePoint, AI-Assisted Development",
            "bullets": [
                "Owned full program lifecycle: stakeholder alignment, requirements, design, development, UAT, launch, and continuous improvement.",
                "Delivered on-time with enterprise-wide adoption within weeks, now a daily-use decision tool for the Change Management team.",
            ],
        },
        {
            "title": "Third-Party Carrier Maintenance Automation",
            "tech": "Python, Outlook COM, Teams Webhooks, openpyxl, Amazon Kiro",
            "bullets": [
                "Drove initiative from concept through production rollout, eliminating a manual process and cutting per-ticket effort by 85%.",
                "Coordinated requirements, built the solution, and managed adoption across the carrier coordination team.",
            ],
        },
        {
            "title": "SOP Standardization Program",
            "tech": "SharePoint, Technical Documentation, Process Design",
            "bullets": [
                "Led the rewrite of 6 SOPs across engineering workflows, aligning execution requirements and reducing onboarding friction for new analysts.",
            ],
        },
    ],
}

# =============================================================================
# HTML GENERATION
# =============================================================================

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>LaKeysha Strickland | {title_line}</title>
    <style>
        :root {{ --navy: #3D4F5F; --accent: #B87333; --dark-text: #2D2D2D; --light-bg: #F5F2EF; --white: #FFFFFF; --border: #E0DEDA; }}
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: var(--dark-text); line-height: 1.6; background: #f4f4f4; }}
        .container {{ max-width: 940px; margin: 30px auto; background: var(--white); box-shadow: 0 4px 20px rgba(0,0,0,0.08); border-radius: 8px; overflow: hidden; }}
        .header {{ background: linear-gradient(135deg, #3D4F5F 0%, #4A6274 100%); color: white; padding: 50px 40px 30px; text-align: center; }}
        .header h1 {{ font-size: 2.4em; font-weight: 700; margin-bottom: 4px; }}
        .header .credentials {{ font-size: 1.1em; color: #D4A574; margin-bottom: 12px; }}
        .header .title-line {{ font-size: 0.85em; font-weight: 600; letter-spacing: 1.5px; text-transform: uppercase; opacity: 0.85; margin-bottom: 16px; }}
        .contact-line {{ font-size: 0.9em; opacity: 0.9; }}
        .contact-line a {{ color: #D4A574; text-decoration: none; }}
        .contact-line .sep {{ margin: 0 8px; opacity: 0.5; }}
        .section {{ padding: 36px 40px; border-bottom: 1px solid var(--border); }}
        .section:last-of-type {{ border-bottom: none; }}
        .section-heading {{ font-size: 1.1em; font-weight: 700; color: var(--navy); text-transform: uppercase; letter-spacing: 1px; border-bottom: 3px solid var(--accent); padding-bottom: 6px; margin-bottom: 18px; }}
        .summary {{ font-size: 0.95em; line-height: 1.8; color: #444; }}
        .skills-category {{ margin-bottom: 16px; }}
        .skills-category-title {{ font-size: 0.85em; font-weight: 700; color: var(--navy); margin-bottom: 8px; text-transform: uppercase; }}
        .skills-list {{ display: flex; flex-wrap: wrap; gap: 8px; }}
        .skill-tag {{ background: var(--light-bg); border: 1px solid var(--border); border-radius: 20px; padding: 6px 14px; font-size: 0.82em; color: var(--navy); font-weight: 500; }}
        .job {{ margin-bottom: 28px; padding-left: 18px; border-left: 3px solid var(--border); }}
        .job-header {{ margin-bottom: 4px; }}
        .job-title {{ font-size: 1.08em; font-weight: 700; color: var(--navy); display: inline; }}
        .job-dates {{ font-size: 0.85em; color: var(--accent); font-weight: 600; float: right; }}
        .job-company {{ font-size: 0.9em; color: #555; margin-bottom: 10px; }}
        .job ul {{ list-style: none; padding: 0; }}
        .job ul li {{ position: relative; padding-left: 18px; margin-bottom: 8px; font-size: 0.9em; line-height: 1.6; }}
        .job ul li::before {{ content: "\\2022"; color: var(--accent); font-weight: bold; position: absolute; left: 0; }}
        .project {{ margin-bottom: 24px; padding-left: 18px; border-left: 3px solid var(--border); }}
        .project-title {{ font-size: 1.05em; font-weight: 700; color: var(--navy); }}
        .project-tech {{ font-size: 0.82em; color: #666; margin-bottom: 8px; font-style: italic; }}
        .project ul {{ list-style: none; padding: 0; }}
        .project ul li {{ position: relative; padding-left: 18px; margin-bottom: 6px; font-size: 0.9em; line-height: 1.6; }}
        .project ul li::before {{ content: "\\2022"; color: var(--accent); font-weight: bold; position: absolute; left: 0; }}
        .education-item {{ margin-bottom: 12px; padding: 12px 16px; background: var(--light-bg); border-radius: 6px; border-left: 3px solid var(--accent); }}
        .education-item .degree {{ font-weight: 700; color: var(--navy); font-size: 0.95em; display: block; }}
        .education-item .school {{ color: #555; font-size: 0.9em; }}
        .education-item .year {{ color: var(--accent); font-weight: 600; font-size: 0.85em; float: right; }}
        .cert-list {{ display: flex; flex-wrap: wrap; gap: 10px; margin: 16px 0; }}
        .cert-tag {{ background: var(--navy); color: white; border-radius: 6px; padding: 10px 18px; font-size: 0.85em; font-weight: 500; }}
        .awards-list {{ list-style: none; padding: 0; }}
        .awards-list li {{ position: relative; padding-left: 18px; margin-bottom: 10px; font-size: 0.92em; line-height: 1.6; }}
        .awards-list li::before {{ content: "\\2022"; color: var(--accent); font-weight: bold; position: absolute; left: 0; }}
        .footer {{ text-align: center; padding: 20px; font-size: 0.8em; color: #999; border-top: 1px solid var(--border); }}
        @media (max-width: 768px) {{ .container {{ margin: 0; border-radius: 0; }} .header {{ padding: 30px 20px; }} .section {{ padding: 24px 20px; }} .job-dates {{ float: none; display: block; }} }}
        @media print {{ body {{ background: white; }} .container {{ box-shadow: none; margin: 0; }} .header {{ background: #3D4F5F !important; -webkit-print-color-adjust: exact; }} .section {{ padding: 20px 30px; }} }}
    </style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>{name}</h1>
        <div class="credentials">{credentials}</div>
        <div class="title-line">{title_line}</div>
        <div class="contact-line">
            {location} <span class="sep">|</span>
            <a href="tel:+1{phone_raw}">{phone}</a> <span class="sep">|</span>
            <a href="mailto:{email}">{email}</a> <span class="sep">|</span>
            <a href="{linkedin}" target="_blank">LinkedIn</a> <span class="sep">|</span>
            <a href="{github}" target="_blank">GitHub</a>
        </div>
    </div>
{body}
    <div class="footer">&copy; 2026 LaKeysha Strickland. Built with Amazon Kiro.</div>
</div>
</body>
</html>"""


def build_html_body(data):
    """Build the HTML body sections from data dict."""
    parts = []

    # Summary
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Professional Summary</h2>')
    parts.append(f'        <p class="summary">{data["summary"]}</p>')
    parts.append(f'    </div>')

    # Skills
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Core Skills</h2>')
    for cat, skills in data["skills"].items():
        parts.append(f'        <div class="skills-category">')
        parts.append(f'            <div class="skills-category-title">{cat}</div>')
        parts.append(f'            <div class="skills-list">')
        for s in skills:
            parts.append(f'                <span class="skill-tag">{s}</span>')
        parts.append(f'            </div>')
        parts.append(f'        </div>')
    parts.append(f'    </div>')

    # Experience
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Professional Experience</h2>')
    for job in data["experience"]:
        parts.append(f'        <div class="job">')
        parts.append(f'            <div class="job-header">')
        parts.append(f'                <div class="job-title">{job["title"]}</div>')
        parts.append(f'                <div class="job-dates">{job["dates"]}</div>')
        parts.append(f'            </div>')
        parts.append(f'            <div class="job-company">{job["company"]}</div>')
        parts.append(f'            <ul>')
        for b in job["bullets"]:
            parts.append(f'                <li>{b}</li>')
        parts.append(f'            </ul>')
        parts.append(f'        </div>')
    parts.append(f'    </div>')

    # Projects
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Key Projects</h2>')
    for proj in data["projects"]:
        parts.append(f'        <div class="project">')
        parts.append(f'            <div class="project-title">{proj["title"]}</div>')
        parts.append(f'            <div class="project-tech">{proj["tech"]}</div>')
        parts.append(f'            <ul>')
        for b in proj["bullets"]:
            parts.append(f'                <li>{b}</li>')
        parts.append(f'            </ul>')
        parts.append(f'        </div>')
    parts.append(f'    </div>')

    # Education & Certs
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Education &amp; Certifications</h2>')
    for degree, school, year in EDUCATION:
        parts.append(f'        <div class="education-item">')
        parts.append(f'            <span class="year">{year}</span>')
        parts.append(f'            <span class="degree">{degree}</span>')
        parts.append(f'            <span class="school">{school}</span>')
        parts.append(f'        </div>')
    parts.append(f'        <div class="cert-list" style="margin-top: 24px;">')
    for cert in CERTIFICATIONS:
        parts.append(f'            <span class="cert-tag">{cert}</span>')
    parts.append(f'        </div>')
    parts.append(f'    </div>')

    # Awards
    parts.append(f'    <div class="section">')
    parts.append(f'        <h2 class="section-heading">Awards &amp; Leadership</h2>')
    parts.append(f'        <ul class="awards-list">')
    for a in AWARDS:
        parts.append(f'            <li>{a}</li>')
    parts.append(f'        </ul>')
    parts.append(f'    </div>')

    return "\n".join(parts)


def generate_html(data, filename):
    """Generate an HTML resume file."""
    body = build_html_body(data)
    html = HTML_TEMPLATE.format(
        name=NAME,
        credentials=CREDENTIALS,
        title_line=data["title_line"],
        location=LOCATION,
        phone=PHONE,
        phone_raw=PHONE.replace("-", ""),
        email=EMAIL,
        linkedin=LINKEDIN,
        github=GITHUB,
        body=body,
    )
    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  [OK] HTML: {filename}")
    return path


# =============================================================================
# PDF GENERATION
# =============================================================================

def sanitize(text):
    replacements = {"\u2013": "-", "\u2014": "--", "\u2018": "'", "\u2019": "'",
                    "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "...", "\u00b7": " - "}
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text


class ResumePDF(FPDF):
    NAVY = (61, 79, 95)
    ACCENT = (184, 115, 51)
    DARK = (45, 45, 45)
    LIGHT = (100, 100, 100)

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def add_header(self, title_line):
        self.set_font("Helvetica", "B", 22)
        self.set_text_color(*self.NAVY)
        self.cell(0, 10, NAME, align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 10)
        self.set_text_color(*self.ACCENT)
        self.cell(0, 6, CREDENTIALS, align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(*self.DARK)
        self.cell(0, 6, sanitize(title_line.upper()), align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*self.DARK)
        contact = f"{LOCATION}  |  {PHONE}  |  {EMAIL}  |  LinkedIn  |  GitHub"
        self.cell(0, 5, contact, align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(6)

    def section_heading(self, text):
        self.ln(3)
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(*self.NAVY)
        self.cell(0, 7, sanitize(text), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*self.ACCENT)
        self.set_line_width(0.5)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body_text(self, text):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*self.DARK)
        self.multi_cell(0, 4.5, sanitize(text))
        self.ln(2)

    def bullet(self, text):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*self.DARK)
        self.set_x(self.get_x() + 4)
        self.multi_cell(0, 4.5, sanitize(f"  -  {text}"))
        self.ln(1)

    def job_head(self, title, company, dates):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*self.NAVY)
        tw = self.get_string_width(sanitize(title))
        self.cell(tw + 2, 6, sanitize(title))
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*self.ACCENT)
        self.cell(0, 6, sanitize(dates), align="R", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*self.LIGHT)
        self.cell(0, 5, sanitize(company), new_x="LMARGIN", new_y="NEXT")
        self.ln(1)


def generate_pdf(data, filename):
    """Generate a PDF resume."""
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 15, 18)
    pdf.add_header(data["title_line"])

    pdf.section_heading("PROFESSIONAL SUMMARY")
    pdf.body_text(data["summary"])

    pdf.section_heading("CORE SKILLS")
    for cat, skills in data["skills"].items():
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(*ResumePDF.NAVY)
        pdf.cell(0, 5, sanitize(cat.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*ResumePDF.DARK)
        pdf.multi_cell(0, 4, sanitize("  |  ".join(skills)))
        pdf.ln(2)

    pdf.section_heading("PROFESSIONAL EXPERIENCE")
    for job in data["experience"]:
        pdf.job_head(job["title"], job["company"], job["dates"])
        for b in job["bullets"]:
            pdf.bullet(b)
        pdf.ln(2)

    pdf.section_heading("KEY PROJECTS")
    for proj in data["projects"]:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*ResumePDF.NAVY)
        pdf.cell(0, 6, sanitize(proj["title"]), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(*ResumePDF.LIGHT)
        pdf.cell(0, 4, sanitize(proj["tech"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
        for b in proj["bullets"]:
            pdf.bullet(b)
        pdf.ln(2)

    pdf.section_heading("EDUCATION & CERTIFICATIONS")
    for degree, school, year in EDUCATION:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*ResumePDF.DARK)
        pdf.cell(0, 5, sanitize(f"{degree}  |  {school}  |  {year}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    for cert in CERTIFICATIONS:
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*ResumePDF.DARK)
        pdf.cell(0, 4.5, sanitize(f"  -  {cert}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    pdf.section_heading("AWARDS & LEADERSHIP")
    for a in AWARDS:
        pdf.bullet(a)

    path = os.path.join(OUTPUT_DIR, filename)
    pdf.output(path)
    print(f"  [OK] PDF:  {filename}")
    return path


# =============================================================================
# WORD GENERATION
# =============================================================================

def generate_word(data, filename):
    """Generate a Word (.docx) resume."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(0)
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x3D, 0x4F, 0x5F)

    # Credentials
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    r = p.add_run(CREDENTIALS)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xB8, 0x73, 0x33)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    r = p.add_run(data["title_line"].upper())
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    r.bold = True

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    contact = f"{LOCATION}  |  {PHONE}  |  {EMAIL}  |  LinkedIn  |  GitHub"
    r = p.add_run(contact)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    def add_heading(text):
        p = doc.add_paragraph()
        p.space_before = Pt(14)
        p.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x3D, 0x4F, 0x5F)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="B87333"/></w:pBdr>')
        pPr.append(pBdr)

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    # Summary
    add_heading("PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    r = p.add_run(data["summary"])
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    # Skills
    add_heading("CORE SKILLS")
    for cat, skills in data["skills"].items():
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        r = p.add_run(f"{cat}: ")
        r.bold = True
        r.font.size = Pt(9)
        r = p.add_run("  |  ".join(skills))
        r.font.size = Pt(9)

    # Experience
    add_heading("PROFESSIONAL EXPERIENCE")
    for job in data["experience"]:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        p.space_after = Pt(0)
        r = p.add_run(job["title"])
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x3D, 0x4F, 0x5F)
        r = p.add_run(f"  |  {job['company']}")
        r.font.size = Pt(9)
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(4)
        r = p2.add_run(job["dates"])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xB8, 0x73, 0x33)
        for b in job["bullets"]:
            add_bullet(b)

    # Projects
    add_heading("KEY PROJECTS")
    for proj in data["projects"]:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        p.space_after = Pt(2)
        r = p.add_run(proj["title"])
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x3D, 0x4F, 0x5F)
        r = p.add_run(f"  |  {proj['tech']}")
        r.font.size = Pt(9)
        for b in proj["bullets"]:
            add_bullet(b)

    # Education
    add_heading("EDUCATION & CERTIFICATIONS")
    for degree, school, year in EDUCATION:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        r = p.add_run(f"{degree}  |  {school}  |  {year}")
        r.bold = True
        r.font.size = Pt(10)
    for cert in CERTIFICATIONS:
        add_bullet(cert)

    # Awards
    add_heading("AWARDS & LEADERSHIP")
    for a in AWARDS:
        add_bullet(a)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  [OK] Word: {filename}")
    return path


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    versions = [
        ("AI Automation Engineer", AI_AUTO, "LaKeysha_Strickland_AI_Automation"),
        ("Solutions Engineer", SOLUTIONS_ENG, "LaKeysha_Strickland_Solutions_Eng"),
        ("Technical Program Manager", TPM, "LaKeysha_Strickland_TPM"),
    ]

    print("=" * 60)
    print("Building Targeted Resumes for LaKeysha Strickland")
    print("=" * 60)

    for label, data, base_name in versions:
        print(f"\n--- {label} ---")
        generate_html(data, f"{base_name}.html")
        generate_pdf(data, f"{base_name}.pdf")
        generate_word(data, f"{base_name}.docx")

    print("\n" + "=" * 60)
    print("Done! All 9 files generated.")
    print("=" * 60)
